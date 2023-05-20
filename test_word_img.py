from docx import Document
from docx.shared import Mm 

doc = Document("sample.docx")

p1 = doc.add_paragraph('画像1を挿入します。')
doc.add_picture('./picturefile/pict1.jpg', width=Mm(50))

count = 0
for para in doc.paragraphs:
    count += len(para.text)
print("このページの文字数は",count)

doc.save("Sample.docx")

doc.save("sample_answer.docx")